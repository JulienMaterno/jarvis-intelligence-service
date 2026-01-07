"""
Document Extractor - Extract text and metadata from various file formats.

Supports:
- PDF: Using PyPDF2 or pdfplumber
- Markdown: Direct text extraction
- TXT: Plain text
- DOCX: Using python-docx (optional)
"""

import logging
import re
from typing import Dict, Any, Optional, Tuple
from pathlib import Path
import io

logger = logging.getLogger("Jarvis.Documents.Extractor")


class DocumentExtractor:
    """Extract text content and metadata from various document formats."""
    
    SUPPORTED_TYPES = {
        "pdf": ["application/pdf"],
        "markdown": ["text/markdown", "text/x-markdown"],
        "text": ["text/plain"],
        "docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
    }
    
    @classmethod
    def extract(
        cls,
        file_bytes: bytes,
        filename: str,
        mime_type: Optional[str] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content and metadata from a document.
        
        Args:
            file_bytes: Raw file bytes
            filename: Original filename
            mime_type: MIME type if known
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        # Determine file type
        ext = Path(filename).suffix.lower().lstrip(".")
        
        if ext == "pdf" or (mime_type and "pdf" in mime_type):
            return cls._extract_pdf(file_bytes, filename)
        elif ext in ("md", "markdown"):
            return cls._extract_markdown(file_bytes, filename)
        elif ext == "txt":
            return cls._extract_text(file_bytes, filename)
        elif ext == "docx":
            return cls._extract_docx(file_bytes, filename)
        else:
            # Try as plain text
            logger.warning(f"Unknown file type: {ext}, trying as text")
            return cls._extract_text(file_bytes, filename)
    
    @classmethod
    def _extract_pdf(cls, file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF."""
        metadata = {"format": "pdf", "filename": filename}
        
        try:
            # Try PyPDF2 first
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
                
                text = "\n\n".join(text_parts)
                metadata["page_count"] = len(reader.pages)
                
                # Extract PDF metadata
                if reader.metadata:
                    if reader.metadata.title:
                        metadata["title"] = reader.metadata.title
                    if reader.metadata.author:
                        metadata["author"] = reader.metadata.author
                    if reader.metadata.creation_date:
                        metadata["created"] = str(reader.metadata.creation_date)
                
                logger.info(f"Extracted {len(text)} chars from PDF: {filename}")
                return text.strip(), metadata
                
            except ImportError:
                logger.warning("PyPDF2 not installed, trying pdfplumber")
                
            # Fallback to pdfplumber
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        text_parts.append(page.extract_text() or "")
                    
                    text = "\n\n".join(text_parts)
                    metadata["page_count"] = len(pdf.pages)
                    
                    logger.info(f"Extracted {len(text)} chars from PDF: {filename}")
                    return text.strip(), metadata
                    
            except ImportError:
                logger.error("Neither PyPDF2 nor pdfplumber installed")
                return "", {"error": "PDF extraction not available", **metadata}
                
        except Exception as e:
            logger.error(f"Failed to extract PDF {filename}: {e}")
            return "", {"error": str(e), **metadata}
    
    @classmethod
    def _extract_markdown(cls, file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Markdown, preserving structure."""
        metadata = {"format": "markdown", "filename": filename}
        
        try:
            text = file_bytes.decode("utf-8")
            
            # Extract title from first # heading
            title_match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
            if title_match:
                metadata["title"] = title_match.group(1).strip()
            
            # Count sections
            headings = re.findall(r"^#{1,3}\s+.+$", text, re.MULTILINE)
            metadata["sections"] = len(headings)
            
            logger.info(f"Extracted {len(text)} chars from Markdown: {filename}")
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Failed to extract Markdown {filename}: {e}")
            return "", {"error": str(e), **metadata}
    
    @classmethod
    def _extract_text(cls, file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract plain text."""
        metadata = {"format": "text", "filename": filename}
        
        try:
            # Try UTF-8 first, then fallback encodings
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    text = file_bytes.decode(encoding)
                    metadata["encoding"] = encoding
                    logger.info(f"Extracted {len(text)} chars from text: {filename}")
                    return text.strip(), metadata
                except UnicodeDecodeError:
                    continue
            
            return "", {"error": "Could not decode text", **metadata}
            
        except Exception as e:
            logger.error(f"Failed to extract text {filename}: {e}")
            return "", {"error": str(e), **metadata}
    
    @classmethod
    def _extract_docx(cls, file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX including tables."""
        metadata = {"format": "docx", "filename": filename}
        
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # IMPORTANT: Also extract text from tables (common in CVs)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n\n".join(text_parts)
            
            # Extract core properties
            if doc.core_properties:
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
            
            logger.info(f"Extracted {len(text)} chars from DOCX: {filename}")
            return text.strip(), metadata
            
        except ImportError:
            logger.error("python-docx not installed")
            return "", {"error": "DOCX extraction not available", **metadata}
        except Exception as e:
            logger.error(f"Failed to extract DOCX {filename}: {e}")
            return "", {"error": str(e), **metadata}
    
    @classmethod
    def infer_document_type(cls, filename: str, content: str) -> str:
        """
        Infer the document type from filename and content.
        
        Returns one of: cv, profile, application, notes, other
        """
        filename_lower = filename.lower()
        content_lower = content.lower()[:2000]  # Check first 2000 chars
        
        # CV detection
        cv_keywords = ["curriculum vitae", "resume", "cv", "work experience", 
                      "education", "skills", "employment history"]
        if any(kw in filename_lower for kw in ["cv", "resume", "lebenslauf"]):
            return "cv"
        if sum(1 for kw in cv_keywords if kw in content_lower) >= 3:
            return "cv"
        
        # Profile detection
        profile_keywords = ["about me", "bio", "introduction", "who i am",
                          "my background", "personal statement"]
        if any(kw in filename_lower for kw in ["profile", "bio", "about"]):
            return "profile"
        if sum(1 for kw in profile_keywords if kw in content_lower) >= 2:
            return "profile"
        
        # Application detection
        app_keywords = ["application", "why i want", "cover letter", 
                       "motivation", "dear hiring", "i am applying"]
        if any(kw in filename_lower for kw in ["application", "cover", "motivation"]):
            return "application"
        if sum(1 for kw in app_keywords if kw in content_lower) >= 2:
            return "application"
        
        # Notes detection
        if any(kw in filename_lower for kw in ["notes", "memo", "thoughts"]):
            return "notes"
        
        return "other"
